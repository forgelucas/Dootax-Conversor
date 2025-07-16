import os
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image

from app.conversor.pptx_utils import extrair_texto_formatado_do_pptx, extrair_texto_imagens_pptx
from app.conversor.notificacoes import notificar_sucesso, notificar_erro

def converter_pptx_para_docx(caminho_pptx):
    try:
        doc = Document()

        # Cria a pasta onde o arquivo final será salvo
        pasta_destino = os.path.join(os.getcwd(), "arquivos_convertidos")
        os.makedirs(pasta_destino, exist_ok=True)

        # Extrai texto e imagens da apresentação
        slides_formatados = extrair_texto_formatado_do_pptx(caminho_pptx)
        imagens_por_slide = extrair_texto_imagens_pptx(caminho_pptx)

        EXTENSOES_VALIDAS = {
            "jpg": "JPEG", "jpeg": "JPEG", "png": "PNG",
            "gif": "GIF", "bmp": "BMP", "tiff": "TIFF", "webp": "WEBP"
        }

        nome_arquivo = os.path.splitext(os.path.basename(caminho_pptx))[0] + ".docx"
        caminho_final = os.path.join(pasta_destino, nome_arquivo)

        for indice, slide in enumerate(slides_formatados):
            for item in slide:
                if item['tipo'] == 'titulo':
                    doc.add_paragraph(item['texto'], style='Heading 2')
                elif item['tipo'] == 'paragrafo':
                    par = doc.add_paragraph()
                    for run in item['runs']:
                        r = par.add_run(run['texto'])
                        if run['bold']:
                            r.bold = True

            # Adiciona imagens diretamente no DOCX (sem salvar)
            if indice < len(imagens_por_slide):
                for img in imagens_por_slide[indice]:
                    try:
                        extensao = img['ext'].lower()
                        formato = EXTENSOES_VALIDAS.get(extensao, "PNG")

                        image = Image.open(BytesIO(img['dados']))
                        buffer = BytesIO()
                        image.save(buffer, format=formato)
                        buffer.seek(0)

                        doc.add_picture(buffer, width=Inches(5))

                    except Exception as e:
                        print(f"Erro ao inserir imagem do slide {indice+1}: {e}")
                        continue

        doc.save(caminho_final)
        notificar_sucesso(caminho_final)

    except Exception as e:
        erro_msg = f"Erro ao converter o arquivo PPTX:\n{caminho_pptx}\n\nDetalhes: {str(e)}"
        print(erro_msg)
        notificar_erro(erro_msg)