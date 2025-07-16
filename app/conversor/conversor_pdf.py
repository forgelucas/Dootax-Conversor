import os
import fitz  
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image

from app.conversor.notificacoes import notificar_sucesso, notificar_erro
from app.conversor.utils import gerar_nome_unico  # opcional, se já estiver usando

EXTENSOES_VALIDAS = {
    "jpg": "JPEG", "jpeg": "JPEG", "png": "PNG",
    "gif": "GIF", "bmp": "BMP", "tiff": "TIFF", "webp": "WEBP"
}

def converter_pdf_para_docx(caminhos):
    pasta_destino = os.path.join(os.getcwd(), "arquivos_convertidos")
    os.makedirs(pasta_destino, exist_ok=True)

    for caminho_pdf in caminhos:
        try:
            doc = Document()
            pdf = fitz.open(caminho_pdf)

            for num_pagina, pagina in enumerate(pdf, start=1):
                doc.add_paragraph(f"Página {num_pagina}:", style='Heading 2')

                texto_simples = pagina.get_text("text").strip()
                doc.add_paragraph(texto_simples if texto_simples else "[Nenhum texto acessível nesta página]")

                imagens = pagina.get_images(full=True)
                for idx_img, img in enumerate(imagens, start=1):
                    xref = img[0]
                    base_img = pdf.extract_image(xref)
                    img_bytes = base_img["image"]
                    ext = base_img["ext"].lower()
                    formato = EXTENSOES_VALIDAS.get(ext, "PNG")

                    try:
                        image = Image.open(BytesIO(img_bytes))
                        buffer = BytesIO()
                        image.save(buffer, format=formato)
                        buffer.seek(0)
                        doc.add_picture(buffer, width=Inches(5))
                    except Exception as erro_img:
                        print(f"Erro ao salvar imagem da página {num_pagina}: {erro_img}")
                        continue

            nome_base = os.path.splitext(os.path.basename(caminho_pdf))[0]
            caminho_docx = os.path.join(pasta_destino, f"{nome_base}.docx")

            if os.path.exists(caminho_docx):
                caminho_docx = gerar_nome_unico(pasta_destino, nome_base)

            doc.save(caminho_docx)
            notificar_sucesso(caminho_docx)

        except Exception as e:
            erro_msg = f"Erro ao converter o arquivo PDF:\n{caminho_pdf}\n\nDetalhes: {str(e)}"
            print(erro_msg)
            notificar_erro(erro_msg)