import os
import xlwings as xw
from docx import Document
from docx.shared import Inches
from app.conversor.notificacoes import notificar_sucesso, notificar_erro
from app.conversor.utils import gerar_nome_unico  

def converter_excel_para_docx(caminhos):
    pasta_destino = os.path.join(os.getcwd(), "arquivos_convertidos")
    os.makedirs(pasta_destino, exist_ok=True)

    for caminho_excel in caminhos:
        try:
            nome_base = os.path.splitext(os.path.basename(caminho_excel))[0]
            caminho_docx = os.path.join(pasta_destino, f"{nome_base}.docx")

            if os.path.exists(caminho_docx):
                caminho_docx = gerar_nome_unico(pasta_destino, nome_base)

            app = xw.App(visible=False)
            wb = app.books.open(caminho_excel)
            doc = Document()

            for planilha in wb.sheets:
                doc.add_paragraph(f"Planilha: {planilha.name}", style='Heading 2')
                dados = planilha.used_range.value

                if not dados or not isinstance(dados, list):
                    continue

                num_cols = max(len(linha) if isinstance(linha, list) else 1 for linha in dados)
                table = doc.add_table(rows=0, cols=num_cols)
                table.style = "Table Grid"

                for linha in dados:
                    row = table.add_row().cells
                    for i in range(num_cols):
                        valor = ""
                        if isinstance(linha, list) and i < len(linha):
                            valor = str(linha[i]) if linha[i] is not None else ""
                        elif not isinstance(linha, list) and i == 0:
                            valor = str(linha)
                        row[i].text = valor

            wb.close()
            app.quit()

            doc.save(caminho_docx)
            notificar_sucesso(caminho_docx)

        except Exception as e:
            erro_msg = f"Erro ao converter Excel:\n{caminho_excel}\n\nDetalhes: {str(e)}"
            print(erro_msg)
            notificar_erro(erro_msg)