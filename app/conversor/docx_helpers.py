import os
from docx import Document
from docx.shared import Inches

def criar_docx_com_formatacao_formatada(slides_formatados, pasta_imagens, caminho_docx):
    doc = Document()
    for slide in slides_formatados:
        slide_num = None
        for bloco in slide:
            if bloco['tipo'] == 'titulo':
                doc.add_paragraph(bloco['texto'])
                try:
                    slide_num = int(bloco['texto'].split()[1].replace(":", ""))
                except:
                    slide_num = None
            elif bloco['tipo'] == 'paragrafo':
                par = doc.add_paragraph(style='List Bullet' if bloco['bullet'] else None)
                for run in bloco['runs']:
                    r = par.add_run(run['texto'])
                    if run['bold']:
                        r.bold = True
        if slide_num is not None:
            for img_file in os.listdir(pasta_imagens):
                if img_file.startswith(f"imagem_{slide_num}_"):
                    caminho_img = os.path.join(pasta_imagens, img_file)
                    doc.add_picture(caminho_img, width=Inches(4.0))
        doc.add_paragraph()
    doc.save(caminho_docx)
